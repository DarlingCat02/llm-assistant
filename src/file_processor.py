import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Максимальная длина текста для передачи в LLM (примерно 6000 символов ≈ 1500 токенов)
MAX_TEXT_LENGTH = 6000


async def extract_text(file_path: str, file_type: str) -> str:
    """
    Извлечь текст из файла.
    
    Args:
        file_path: Путь к файлу
        file_type: Тип файла (txt, pdf, docx)
    
    Returns:
        str: Извлечённый текст
    """
    try:
        ext = file_type.lower().replace("application/", "").replace("text/", "")
        
        if ext == "txt" or ext == "plain":
            return await _extract_txt(file_path)
        elif ext == "pdf":
            return await _extract_pdf(file_path)
        elif ext in ("docx", "vnd.openxmlformats-officedocument.wordprocessingml.document"):
            return await _extract_docx(file_path)
        else:
            return f"[Формат {ext} не поддерживается]"
            
    except Exception as e:
        logger.error(f"Text extraction error from {file_type}: {e}")
        return f"[Ошибка извлечения текста: {e}]"


async def _extract_txt(file_path: str) -> str:
    """Извлечь текст из TXT файла."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="cp1251") as f:
            text = f.read()
    
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + f"\n\n[... обрезано, всего символов: {len(text)}]"
    
    return text


async def _extract_pdf(file_path: str) -> str:
    """Извлечь текст из PDF файла."""
    import pdfplumber
    
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages[:20]):  # Максимум 20 страниц
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Страница {i + 1}]\n{page_text}")
    
    text = "\n\n".join(text_parts)
    
    if not text.strip():
        return "[PDF не содержит извлекаемого текста]"
    
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + f"\n\n[... обрезано, всего символов: {len(text)}]"
    
    return text


async def _extract_docx(file_path: str) -> str:
    """Извлечь текст из DOCX файла."""
    from docx import Document
    
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    
    text = "\n\n".join(text_parts)
    
    if not text.strip():
        return "[DOCX не содержит текста]"
    
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + f"\n\n[... обрезано, всего символов: {len(text)}]"
    
    return text
